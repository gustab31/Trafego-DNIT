# =====================================================
# imports

from pathlib import Path
import glob
import sqlite3 as lite
import sys
import itertools
import pandas as pd
import numpy as np
import csv
from csv import reader
from tkinter import *
import math
import statistics as sta 
import matplotlib.pyplot as plt
import scipy.stats as st
from scipy.stats import norm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document 
import os
import win32com.client
import statsmodels.formula.api as smf
import comtypes.client
import subprocess
from mpl_toolkits.mplot3d import Axes3D
import random 
import seaborn as sns
import ijson

# =====================================================
# Dados de entrada
input_file_path = "C:\\Users\\gusta\Dropbox\\PC\\Documents\\mestrado\\programacao_trafego\\dados\\origem_destino2.csv"
arquivo_docx = "C:\\Users\\gusta\Dropbox\\PC\\Documents\\mestrado\\programacao_trafego\\resultados\\trabalho_relatorio.docx"
pdf_path = "C:\\Users\\gusta\Dropbox\\PC\\Documents\\mestrado\\programacao_trafego\\resultados\\relatorio.pdf"
input_excel = "C:\\Users\\gusta\Dropbox\\PC\\Documents\\mestrado\\programacao_trafego\\dados\\origem_destino2.xlsx"
saida_csv = "C:\\Users\\gusta\Dropbox\\PC\\Documents\\mestrado\\programacao_trafego\\dados\\tme.csv"
# =====================================================
# Lendo o PDF com os dados
data = pd.read_csv(input_file_path, sep=';', header=None)
data_frame = pd.read_excel(input_excel)
#print (data)
# =====================================================
# Colocando os dados em uma tabela
def df_to_docx_table( dataframe ):

    n_linhas, n_colunas = dataframe.shape
    table = document.add_table(rows = 1, cols = n_colunas, style = "Table Grid")

    # Adicionar nomes das colunas
    hdr_cells = table.rows[0].cells
    colunas = dataframe.columns
    for i, column in enumerate( colunas ):
        hdr_cells[i].text = column 

    # Colocar o conteúdo de dados
    for n_linha in range ( n_linhas ): 
        row_cells = table.add_row().cells
        for n_coluna, cell in  enumerate( row_cells ):
            cell.text = "%d" % int( dataframe.iloc[n_linha][n_coluna])
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# =====================================================

# =====================================================
# Procedimento
document = Document()
document.add_heading('Relatório Tempo de Espera para Acessar a Rotúla',0)

document.add_heading('Fluxo de Veículos', level=1)

                                
p = document.add_paragraph('Dados')
p.add_run('bold').bold = True
df_to_docx_table( data_frame )


document.add_page_break()

# =====================================================
# constantes
tg = 4.1 # valor médio do intervalo mínimo entre veículos na rotatória, segundos
tf = 2.9 # valor médio do intervalo entre dois veículos sucessivos na entrada, segundos
tmin = 2.1 # valor mínimo do intervalo entre veículos da rotatória, segundos
nk = 2 # número de faixas de tráfego na pista rotatória antes da entrada i
nz = 2 # número de faixas de tráfego na entrada i

# =====================================================
#trabalhando com os dados
separador = ';'

z1 = []
z2 = []
z3 = []
z4 = []
k1 = []
k2 = []
k3 = []
k4 = []
with open(input_file_path, 'r', newline='') as csv_file:
    for line_number, content in enumerate(csv_file):
        if line_number:  # pula cabeçalho
            colunas = content.strip().split(separador)
            z1.append( float(colunas[0]) )
            z2.append( float(colunas[1]) )
            z3.append( float(colunas[2]) )
            z4.append( float(colunas[3]) )
            k1.append( float(colunas[4]) )
            k2.append( float(colunas[5]) )
            k3.append( float(colunas[6]) )
            k4.append( float(colunas[7]) )
            # avaliando a condicao de trafego no futuro

# Página 197 manual de tráfego DNIT (2006)
# =====================================================
f= open(saida_csv, 'w', newline='', encoding='utf-8') 
w = csv.writer(f,delimiter="\n") 

# =====================================================
# Capacidade de trafego em cada entrada

 
Ri = []   
TME = []
Ci = []
dados = []
df = []
n = 0 
lista_Ci = []  
lista_Ri = []
lista_Zi = []
lista_TME = []
tempo_espera1 = []
rows = []
def volume_trafego1(input1,input2):  
    for n in range(len(z1)):


    # Página 197 manual de tráfego DNIT (2006)
      
    # Os volumes de tráfego na rotatória, em UCP/h         
        ki = input1[n]
                    
    # Fluxo de entrada, em UCP/h
                    
        Zi = input2[n]

    # Fator de pedestre
        fi = 0.152*(ki**(0.1701)) # fator de pedestre tabelado
                  
    # Capacidade básica da entrada, em UCP/h
                             
        Gi = ((3600*( 1 - ((tmin*ki/(nk*3600)))**nk)*nz/tf))*math.exp((-ki/3600*(tg - tf/2 - tmin)))
        #print(Gi)
                    
    # Capacidade de entrada
        Ci = Gi*fi
        #print(Ci)
                    
    # Capacidade Residual
        Ri = Ci - Zi
        
        #Ri.append( [float(Ri)] )
        
    # Tempo médio de espera (segundos)
        TME = 195.612*(Ri**(-0.517))
                      
        dados = {'Capacidade Entrada (UCP/h)': [Ci], 'Capacidade Residual (UCP/h)':[Ri],'Tempo Médio de Espera (s)':[TME]}
        df = pd.DataFrame(dados)
        print(df)
                    
        
             
        document.add_heading("Avaliação do Volume de Tráfego em UCP/h", 1)
        df_to_docx_table(df)

        #print(TME)  
        
        

        lista_Ci.append(int(Ci))
        lista_Ri.append(int(Ri))
        lista_TME.append(int(TME))
    dados2 = [lista_Ci,lista_Ri,lista_TME]
    df2 = pd.DataFrame(data=dados2)
# =====================================================         
       
    
    w.writerow([df2]) 
    #w.writerow(lista_Ci) 
          
        
   
# =====================================================       

tempo_espera1 = volume_trafego1(k1,z1)
#tempo_espera2 = volume_trafego2(k2,z2)
#tempo_espera3 = volume_trafego3(k3,z3)
#tempo_espera4 = volume_trafego4(k4,z4)

# =====================================================   
document.add_page_break()
  
# =====================================================                                 
# Relatorio 

document.save( arquivo_docx )

#Convertendo para PDF

wdFormatPDF = 17

in_file = os.path.abspath(arquivo_docx)
out_file = os.path.abspath(pdf_path)

word = win32com.client.Dispatch('Word.Application')
doc = word.Documents.Open(in_file)
doc.SaveAs(out_file, FileFormat=wdFormatPDF)
doc.Close()
word.Quit()


# =====================================================                

# ===================================================== 
print ("Fim")              
