# =====================================================
# imports

from pathlib import Path
import glob
import sqlite3 as lite
import sys
import itertools
import pandas as pd
import numpy as np
import csv
from csv import reader
from tkinter import *
import math
import statistics as sta 
import matplotlib.pyplot as plt
import scipy.stats as st
from scipy.stats import norm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document 
import os
import win32com.client
import statsmodels.formula.api as smf
import comtypes.client
import subprocess
from mpl_toolkits.mplot3d import Axes3D
import random 
import seaborn as sns
# =====================================================
# Dados de entrada
input_file_path = "C:\\Users\\gusta\Dropbox\\PC\\Documents\\mestrado\\programacao_trafego\\dados\\time.csv"
arquivo_docx = "C:\\Users\\gusta\Dropbox\\PC\\Documents\\mestrado\\programacao_trafego\\resultados\\trabalho_tme.docx"
pdf_path = "C:\\Users\\gusta\Dropbox\\PC\\Documents\\mestrado\\programacao_trafego\\resultados\\relatorio_tme.pdf"
input_excel = "C:\\Users\\gusta\Dropbox\\PC\\Documents\\mestrado\\programacao_trafego\\dados\\time.xlsx"
# =====================================================
# Lendo o PDF com os dados
data = pd.read_csv(input_file_path, sep=';', header=None)
data_frame = pd.read_excel(input_excel)
#print (data)
# =====================================================
# Colocando os dados em uma tabela
def df_to_docx_table( dataframe ):

    n_linhas, n_colunas = dataframe.shape
    table = document.add_table(rows = 1, cols = n_colunas, style = "Table Grid")

    # Adicionar nomes das colunas
    hdr_cells = table.rows[0].cells
    colunas = dataframe.columns
    for i, column in enumerate( colunas ):
        hdr_cells[i].text = column 

    # Colocar o conteúdo de dados
    for n_linha in range ( n_linhas ): 
        row_cells = table.add_row().cells
        for n_coluna, cell in  enumerate( row_cells ):
            cell.text = "%d" % int( dataframe.iloc[n_linha][n_coluna])
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# =====================================================
# Procedimento
document = Document()
document.add_heading('Relatório Tempo de Espera para Acessar a Rotúla',0)

document.add_heading('Fluxo de Veículos', level=1)

                                
p = document.add_paragraph('Dados')
p.add_run('bold').bold = True
df_to_docx_table( data_frame )

document.add_page_break()

# =====================================================
#trabalhando com os dados
separador = ';'

z1 = []
z2 = []
z3 = []
z4 = []
TME1 = []
TME2 = []
TME3 = []
TME4 = []
with open(input_file_path, 'r', newline='') as csv_file:
    for line_number, content in enumerate(csv_file):
        if line_number:  # pula cabeçalho
            colunas = content.strip().split(separador)
            z1.append( float(colunas[0]) )
            z2.append( float(colunas[1]) )
            z3.append( float(colunas[2]) )
            z4.append( float(colunas[3]) )
            TME1.append( float(colunas[4]) )
            TME2.append( float(colunas[5]) )
            TME3.append( float(colunas[6]) )
            TME4.append( float(colunas[7]) )
            # avaliando o nivel de serventia viario
            
z1_soma = (sum([i for i in z1 if isinstance(i, int) or isinstance(i, float)]))
z2_soma = (sum([i for i in z2 if isinstance(i, int) or isinstance(i, float)]))           
z3_soma = (sum([i for i in z3 if isinstance(i, int) or isinstance(i, float)]))
z4_soma = (sum([i for i in z4 if isinstance(i, int) or isinstance(i, float)]))
fluxo_tempo1 = (sum([i for i in TME1 if isinstance(i, int) or isinstance(i, float)]))
fluxo_tempo2 = (sum([i for i in TME2 if isinstance(i, int) or isinstance(i, float)]))
fluxo_tempo3 = (sum([i for i in TME3 if isinstance(i, int) or isinstance(i, float)]))
fluxo_tempo4 = (sum([i for i in TME4 if isinstance(i, int) or isinstance(i, float)]))
# ==========================
# Página 203 manual de tráfego DNIT (2006)
# Capacidade na entrada 1

# =====================================================
def nivel_servico(input1):  
    if input1 > 0 and input1 <= 10:
        return "A"
    elif input1 > 10 and input1 <= 20:
        return "B"
    elif input1 > 20 and input1 <= 30:
        return "C"
    elif input1 > 30 and input1 <= 45:
        return "D"
    elif input1 > 45:
        return "E"
    else:
        return "F" 
# =====================================================
fluxo_tempo1 = z1_soma*fluxo_tempo1
fluxo_tempo2 = z2_soma*fluxo_tempo2
fluxo_tempo3 = z3_soma*fluxo_tempo3
fluxo_tempo4 = z4_soma*fluxo_tempo4
#fluxo_tempo = fluxo_tempo1+fluxo_tempo2+fluxo_tempo3+fluxo_tempo4
soma_z = z1_soma+z2_soma+z3_soma+z4_soma
TIMER_1 = fluxo_tempo1/soma_z 
TIMER_2 = fluxo_tempo2/soma_z 
TIMER_3 = fluxo_tempo3/soma_z 
TIMER_4 = fluxo_tempo4/soma_z 
print(TIMER_1)
print(TIMER_2)
print(TIMER_3)
print(TIMER_4)
TIMER=(TIMER_1+TIMER_2+TIMER_3+TIMER_4)/4
LOS = nivel_servico(TIMER)

print("Nível de serviço = %s" % LOS) 
  
 # =====================================================   
document.add_page_break()
  
# =====================================================                                 
# Relatorio 

document.save( arquivo_docx )

#Convertendo para PDF

wdFormatPDF = 17

in_file = os.path.abspath(arquivo_docx)
out_file = os.path.abspath(pdf_path)

word = win32com.client.Dispatch('Word.Application')
doc = word.Documents.Open(in_file)
doc.SaveAs(out_file, FileFormat=wdFormatPDF)
doc.Close()
word.Quit()
