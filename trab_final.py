# Método de Monte Carlo aplicado no tempo de espera para acessar uma rotátoria
# =====================================================
#imports 
from pathlib import Path
import glob
import sqlite3 as lite
import sys
import itertools
import pandas as pd
import numpy as np
import csv
from csv import reader
from tkinter import *
import math
import statistics as sta 
import matplotlib.pyplot as plt
import scipy.stats as st
from scipy.stats import norm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document 
import os
import win32com.client
import statsmodels.formula.api as smf
import comtypes.client
import subprocess
from mpl_toolkits.mplot3d import Axes3D
import random 
import seaborn as sns

# =====================================================
# Dados de entrada
input_file_path = "C:\\Users\\Sucesso\\Documents\\Mestrado Engenharia\\programacao_comp_transp\\trab_final3\\dados\\dados_origem_destino.csv"
arquivo_docx = "C:\\Users\\Sucesso\\Documents\\Mestrado Engenharia\\programacao_comp_transp\\trab_final3\\resultados\\trabalho_relatorio.docx"
pdf_path = "C:\\Users\\Sucesso\\Documents\\Mestrado Engenharia\\programacao_comp_transp\\trab_final3\\resultados\\relatorio.pdf"
input_excel = "C:\\Users\\Sucesso\\Documents\\Mestrado Engenharia\\programacao_comp_transp\\trab_final3\\dados\\dados_origem_destino_int.xlsx"
# =====================================================
# Lendo o PDF com os dados
data = pd.read_csv(input_file_path)
data_frame = pd.read_excel(input_excel)
#print (data)

# =====================================================

# constantes
tg = 4.1 # valor médio do intervalo mínimo entre veículos na rotatória, segundos
tf = 2.9 # valor médio do intervalo entre dois veículos sucessivos na entrada, segundos
tmin = 2.1 # valor mínimo do intervalo entre veículos da rotatória, segundos
nk = 2 # número de faixas de tráfego na pista rotatória antes da entrada i
nz = 2 # número de faixas de tráfego na entrada i
fi = 1 # fator de pedestre
# =====================================================
# Colocando os dados em uma tabela
def df_to_docx_table( dataframe ):

    n_linhas, n_colunas = dataframe.shape
    table = document.add_table(rows = 1, cols = n_colunas, style = "Table Grid")

    # Adicionar nomes das colunas
    hdr_cells = table.rows[0].cells
    colunas = dataframe.columns
    for i, column in enumerate( colunas ):
        hdr_cells[i].text = column 

    # Colocar o conteúdo de dados
    for n_linha in range ( n_linhas ): 
        row_cells = table.add_row().cells
        for n_coluna, cell in  enumerate( row_cells ):
            cell.text = "%d" % int( dataframe.iloc[n_linha][n_coluna])
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# =====================================================
# Procedimento
document = Document()
document.add_heading('Relatório Tempo de Espera para Acessar a Rotúla',0)

document.add_heading('Fluxo de Veículos', level=1)

                                
p = document.add_paragraph('Dados')
p.add_run('bold').bold = True
df_to_docx_table( data_frame )

document.add_page_break()

# =====================================================
#trabalhando com os dados
separador = ';'

lista_k = []
lista_Z = []

with open(input_file_path, 'r', newline='') as csv_file:
    for line_number, content in enumerate(csv_file):
        if line_number:  # pula cabeçalho
            colunas = content.strip().split(separador)
            lista_k.append( float(colunas[5]) )
            lista_Z.append( float(colunas[6]) )
            # avaliando a condicao de trafego no futuro

#print( lista_k )
#print( max(lista_k) )
#print( lista_Z )
#print( max(lista_Z) )

i = 0 
n = 0
ano = 2012 # inicio  
lista_Ri = []        
for n in range(len(lista_k)):
    i += 3 
    taxa_frota = (1+i/100)**n
    k = lista_k[n]
    Z = lista_Z[n]
    #print( k,Z )
    
    #n += 1

# Página 197 manual de tráfego DNIT (2006)
  
# Os volumes de tráfego na rotatória, em UCP/h         
    ki = k*taxa_frota
                
# Fluxo de entrada, em UCP/h
                
    Zi = Z*taxa_frota
                
# Capacidade básica da entrada, em UCP/h
                         
    Gi = ((3600*( 1 - (tmin*ki/(nk*3600)))*nk*nz/tf))*math.exp((-ki/3600*(tg - tf/2 - tmin)))
    #print(Gi)
                
# Capacidade de entrada
    Ci = Gi*fi
    #print(Ci)
                
# Capacidade Residual
    Ri = Ci - Zi
    
    lista_Ri.append( float(Ri) )
                  
    ano += 1 
    #print(ano)
    dados = {'Ano' :[ano], 'Capacidade Entrada': [Ci], 'Capacidade Residual':[Ri]}
    df = pd.DataFrame(dados)
    #print(df)
                
    df2 = pd.DataFrame({'Capacidade Residual':[Ri]})
       
    
    
    document.add_heading("Avaliação do Volume de Tráfego em UCP/h", 1)
    df_to_docx_table(df)
 
Ri_min = min(lista_Ri)
Ri_max = max(lista_Ri)
#print(Ri_min)
#print(Ri_max)

 # =====================================================   
document.add_page_break()

 # =====================================================   
# Valores de Ri aleteatórios 

s = np.random.normal(Ri_min,Ri_max,100)
#print(s)

variancia_s = np.var(s)
moda = (sum(s)/100) 
#print(np.std(s))
# =====================================================  

# Gráfico da Distribuição normal para os valores aleatórios
plt.hist(s, density=True, bins=30, label="Dados") #bins = 30
mn, mx = plt.xlim()
plt.xlim(mn, mx)
kde_xs = np.linspace(mn, mx, 301) #estava 301
kde = st.gaussian_kde(s)
plt.plot(kde_xs, kde.pdf(kde_xs))  #kde.pdf mudar para uma funcao normal 
plt.legend(loc="upper left")
plt.ylabel('Probabilidade')
plt.xlabel('Capacidade Residual')
plt.title("Histograma")
#plt.show("Histogram")
plt.savefig("Histograma.png")         
                
plt.clf()

 # =====================================================   
# Monte Carlo

ultimo = s[-1]
amostras = 1000
simulacao_df = pd.DataFrame()
for a in range(amostras):
    count = 0
    variancia = s.std() # variancia
    
    valor_series = []
    R2 = ultimo*(1+np.random.normal(0,variancia))
    valor_series.append(R2)
    
        
    simulacao_df[a] = valor_series
#print(np.std(simulacao_df[a]))

 # =====================================================         
# Gráfico da Distribuição normal
x = simulacao_df
plt.hist(x, density=True, bins=200, label="Dados") #bins = 30
mn, mx = plt.xlim()
plt.xlim(mn, mx)
kde_xs = np.linspace(mn, mx, 1000) #tava 301
kde = st.gaussian_kde(x)
plt.plot(kde_xs, kde.pdf(kde_xs))  #kde.pdf mudar para uma funcao normal 
plt.legend(loc="upper left")
plt.ylabel('Probabilidade')
plt.xlabel('Capacidade Residual')
plt.title("Histograma")
#plt.show("Histogram")
plt.savefig("Histograma_normal.png")         
                
plt.clf()

 # =====================================================               
                
                
document.add_heading("Histograma", 1)
figura1 = document.add_picture('Histograma.png')

document.add_page_break()
                
document.add_heading("Histograma Monte Carlo", 1)
figura1 = document.add_picture('Histograma_normal.png')

# =====================================================                                 
# Relatorio 

document.save( arquivo_docx )

#Convertendo para PDF

wdFormatPDF = 17

in_file = os.path.abspath(arquivo_docx)
out_file = os.path.abspath(pdf_path)

word = win32com.client.Dispatch('Word.Application')
doc = word.Documents.Open(in_file)
doc.SaveAs(out_file, FileFormat=wdFormatPDF)
doc.Close()
word.Quit()

                


print ("Fim")              

                
                
                
       
